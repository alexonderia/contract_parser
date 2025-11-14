"""Utilities for exporting specification tables into standalone DOCX files."""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document

from .schemas import SpecificationResponse, SpecificationTable

_DEFAULT_EXPORT_DIR = Path(__file__).resolve().parent / "exports"


def _sanitize_stem(value: str) -> str:
    """Return a filesystem-safe stem for the generated document."""

    sanitized = re.sub(r"[^0-9A-Za-zА-Яа-я_-]+", "_", value).strip("._")
    return sanitized or "specification"


def _ensure_export_dir(path: Path | None) -> Path:
    export_dir = Path(path or _DEFAULT_EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def _remove_placeholder_paragraph(document: Document) -> None:
    """Remove the placeholder paragraph that python-docx creates by default."""

    if document.paragraphs:
        paragraph = document.paragraphs[0]
        element = paragraph._element  # type: ignore[attr-defined]
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _table_column_count(rows: Iterable[list[str]]) -> int:
    return max((len(row) for row in rows), default=0)


def _append_table(document: Document, table: SpecificationTable) -> None:
    rows = table.rows or []
    if not rows:
        return

    column_count = _table_column_count(rows)
    if column_count == 0:
        return

    docx_table = document.add_table(rows=len(rows), cols=column_count)
    docx_table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            docx_table.cell(row_index, column_index).text = value or ""


def _pick_filename(source_name: str | None, stem_fallback: str) -> str:
    stem = Path(source_name or "").stem or stem_fallback
    sanitized = _sanitize_stem(stem)
    return f"{sanitized}_specification.docx"


def _next_available_path(directory: Path, filename: str) -> Path:
    candidate = directory / filename
    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix
    counter = 1
    while True:
        updated = directory / f"{stem}_{counter}{suffix}"
        if not updated.exists():
            return updated
        counter += 1


def export_specification_to_docx(
    specification: SpecificationResponse,
    *,
    source_filename: str | None = None,
    export_dir: Path | None = None,
) -> tuple[Path, bytes] | None:
    """Create a DOCX file containing only the specification tables."""

    tables = [table for table in specification.tables if table.rows]
    if not tables:
        return None

    document = Document()
    _remove_placeholder_paragraph(document)

    for table in tables:
        _append_table(document, table)

    buffer = BytesIO()
    document.save(buffer)
    payload = buffer.getvalue()

    export_directory = _ensure_export_dir(export_dir)
    filename = _pick_filename(source_filename, specification.heading or "specification")
    target_path = _next_available_path(export_directory, filename)
    target_path.write_bytes(payload)

    return target_path, payload


__all__ = ["export_specification_to_docx"]